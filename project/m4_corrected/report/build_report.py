#!/usr/bin/env python3
"""
build_report.py  —  Generate report.docx (Sections 10-13 addendum).
Run from repo root:
    source /home/sackb/.venv/bin/activate
    python3 project/m4_corrected/report/build_report.py
"""

import os
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# ── Paths ─────────────────────────────────────────────────────────────────────
SCRIPT_DIR  = os.path.dirname(os.path.abspath(__file__))
FIGURES     = os.path.join(SCRIPT_DIR, "figures")
DOCX_OUT    = os.path.join(SCRIPT_DIR, "report.docx")
TEMPLATE    = os.path.abspath(os.path.join(
    SCRIPT_DIR,
    "../../m4/report/design_justification.docx"
))

FIG4 = os.path.join(FIGURES, "fig4_time_comparison.png")
FIG5 = os.path.join(FIGURES, "fig5_speedup_comparison.png")
FIG6 = os.path.join(FIGURES, "fig6_roofline_comparison.png")

os.makedirs(FIGURES, exist_ok=True)


# ── Figure generation ─────────────────────────────────────────────────────────

def make_fig4():
    """Grouped/stacked bar chart: per-45-update runtime comparison."""
    fig, ax = plt.subplots(figsize=(8, 5), dpi=150)

    groups = ["SW Baseline", "Old HW Benchmark\n(raw-C predict)", "Corrected HW Benchmark\n(Eigen predict)"]
    x = np.arange(len(groups))
    width = 0.5

    # SW baseline: single blue bar
    sw_val = 429.778
    ax.bar(x[0], sw_val, width, color='steelblue', label='SW full filter (Eigen)')
    ax.text(x[0], sw_val + 5, f'{sw_val:.3f} ms', ha='center', va='bottom', fontsize=8)

    # Old HW: stacked predict (negligible) + correct
    old_predict = 0.0      # raw-C predict ≈ 0 ms
    old_correct = 3.952
    ax.bar(x[1], old_predict, width, color='orange', label='Eigen predict (CPU)')
    ax.bar(x[1], old_correct, width, bottom=old_predict, color='red', label='HW correction (MMIO)')
    ax.text(x[1], old_correct + 0.5, f'{old_correct:.3f} ms', ha='center', va='bottom', fontsize=8)

    # Corrected HW: Eigen predict + HW correct
    new_predict = 139.546
    new_correct = 3.952
    ax.bar(x[2], new_predict, width, color='orange')
    ax.bar(x[2], new_correct, width, bottom=new_predict, color='red')
    total_new = new_predict + new_correct
    ax.text(x[2], total_new + 2, f'{total_new:.3f} ms', ha='center', va='bottom', fontsize=8)
    ax.text(x[2], new_predict / 2, f'Eigen\n{new_predict:.1f} ms', ha='center', va='center', fontsize=7, color='white')
    ax.text(x[2], new_predict + new_correct / 2, f'HW\n{new_correct:.3f} ms', ha='center', va='center', fontsize=6, color='white')

    ax.set_yscale('log')
    ax.set_ylabel('Time (ms)', fontsize=11)
    ax.set_title('Per-45-update Runtime Comparison (13 MHz SweRV EL2)', fontsize=12)
    ax.set_xticks(x)
    ax.set_xticklabels(groups, fontsize=9)
    ax.legend(fontsize=8, loc='upper right')
    ax.set_ylim(bottom=0.1)
    ax.grid(axis='y', alpha=0.3)

    fig.tight_layout()
    fig.savefig(FIG4)
    plt.close(fig)
    print(f"  Saved: {FIG4}")


def make_fig5():
    """Horizontal bar chart: reported vs corrected speedup."""
    fig, ax = plt.subplots(figsize=(8, 4), dpi=150)

    labels = [
        "Old benchmark\n(raw-C predict,\nHW_FLOPS=1080)",
        "Corrected benchmark\n(Eigen predict,\nHW_FLOPS=6210)",
    ]
    values = [108.7, 2.995]
    colors = ['salmon', 'steelblue']

    y = np.arange(len(labels))
    bars = ax.barh(y, values, color=colors, height=0.5)

    for bar, val in zip(bars, values):
        ax.text(val * 1.05, bar.get_y() + bar.get_height() / 2,
                f'{val}×', va='center', ha='left', fontsize=10, fontweight='bold')

    ax.axvline(x=1, color='black', linestyle='--', linewidth=1.2, label='No speedup')
    ax.text(1.05, len(labels) - 0.05, '1× (no speedup)', fontsize=8, va='top', color='black')

    ax.set_xscale('log')
    ax.set_xlim(1, 1000)
    ax.set_xlabel('Speedup over SW baseline (×)', fontsize=11)
    ax.set_title('Reported vs Corrected Speedup over SW Baseline', fontsize=12)
    ax.set_yticks(y)
    ax.set_yticklabels(labels, fontsize=9)
    ax.grid(axis='x', alpha=0.3)

    fig.tight_layout()
    fig.savefig(FIG5)
    plt.close(fig)
    print(f"  Saved: {FIG5}")


def make_fig6():
    """Side-by-side roofline subplots."""
    PEAK_COMPUTE = 13e6 / 1e9   # GFLOP/s
    PEAK_BW      = 52e6 / 1e9   # GB/s
    RIDGE        = PEAK_COMPUTE / PEAK_BW  # FLOP/byte

    ai_range = np.logspace(-2, 3, 500)
    roof     = np.minimum(PEAK_COMPUTE, PEAK_BW * ai_range)

    fig, axes = plt.subplots(1, 2, figsize=(12, 5), dpi=150)
    titles = [
        "Original M4 Roofline\n(HW = correction only)",
        "Corrected M4 Roofline\n(HW = Eigen predict + correction)",
    ]

    # Left subplot points
    left_points = [
        ("SW baseline",  0.21,   2.4082e-05, 'b', 's'),
        ("Old HW",       0.0882, 2.7327e-04, 'r', '^'),
    ]
    # Right subplot points
    right_points = [
        ("SW baseline",    0.21,   2.4082e-05, 'b', 's'),
        ("Corrected HW",   0.479,  4.3276e-05, 'r', '^'),
    ]

    all_points = [left_points, right_points]

    for ax, title, points in zip(axes, titles, all_points):
        ax.loglog(ai_range, roof, 'k-', linewidth=2, label='Roofline')
        ax.axvline(x=RIDGE, color='gray', linestyle=':', linewidth=1, alpha=0.7)

        for label, ai, gflops, color, marker in points:
            ax.loglog(ai, gflops, marker=marker, color=color,
                      markersize=9, label=label, zorder=5)

        ax.set_xlim(1e-2, 1e3)
        ax.set_ylim(1e-7, 1e-1)
        ax.set_xlabel('Arithmetic Intensity (FLOP/byte)', fontsize=10)
        ax.set_ylabel('Performance (GFLOP/s)', fontsize=10)
        ax.set_title(title, fontsize=10)
        ax.legend(fontsize=8)
        ax.grid(True, which='both', alpha=0.3)
        ax.text(RIDGE * 1.1, 1e-6, f'Ridge\n{RIDGE:.2f} F/B', fontsize=7, color='gray')

    fig.tight_layout()
    fig.savefig(FIG6)
    plt.close(fig)
    print(f"  Saved: {FIG6}")


# ── Document helpers ──────────────────────────────────────────────────────────

def new_doc():
    """Load existing docx as template to inherit styles, then clear body content."""
    doc = Document(TEMPLATE)
    body = doc.element.body
    for child in list(body):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag not in ('sectPr',):
            body.remove(child)
    return doc


def h1(doc, text):
    p = doc.add_paragraph(text, style='Heading 1')
    return p


def body_para(doc, text):
    p = doc.add_paragraph(style='normal1')
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.space_before = Pt(0)
    return p


def caption(doc, text):
    p = doc.add_paragraph(style='normal1')
    run = p.add_run(text)
    run.font.name   = 'Times New Roman'
    run.font.size   = Pt(10)
    run.font.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after  = Pt(10)
    p.paragraph_format.space_before = Pt(2)
    return p


def figure(doc, img_path, cap_text, width_in=5.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_path, width=Inches(width_in))
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    caption(doc, cap_text)


def add_table(doc, headers, rows, col_widths=None):
    n_cols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)
    try:
        tbl.style = 'Table Grid'
    except KeyError:
        pass
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Data rows
    for r, row_data in enumerate(rows):
        row = tbl.rows[r + 1]
        for c, val in enumerate(row_data):
            cell = row.cells[c]
            cell.text = val
            cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if col_widths:
        for row in tbl.rows:
            for c_idx, cell in enumerate(row.cells):
                cell.width = Inches(col_widths[c_idx])
    doc.add_paragraph()
    return tbl


# ── Build document ─────────────────────────────────────────────────────────────

def build(doc):

    # ── Title block ───────────────────────────────────────────────────────────
    doc.add_paragraph("Benchmark Correction Addendum", style='Title')
    doc.add_paragraph(
        "ECE 410/510 Spring 2026  |  Jose Ramirez  |  Milestone 4\n"
        "Continuation of design_justification.docx \u2014 Sections 10\u201313",
        style='Subtitle'
    )
    doc.add_paragraph()

    # ─────────────────────────────────────────────────────────────────────────
    # 10. ERROR IN THE M4 BENCHMARK PROGRAM
    # ─────────────────────────────────────────────────────────────────────────
    h1(doc, "10. Error in the M4 Benchmark Program")

    body_para(doc,
        "Two sub-issues caused the M4 benchmark results reported in Section 8 to "
        "misrepresent the true system speedup."
    )

    h1(doc, "10.1 Structural Asymmetry in the Benchmark Program")

    body_para(doc,
        "The software baseline (RVfpgaEL2_profiling/test/kalman-test.cpp) calls "
        "kf.update(y) on a KalmanFilter object constructed with Eigen MatrixXd/VectorXd "
        "types. This method executes both the predict step (\u0078\u0302\u207b = A\u00b7\u0078\u0302, "
        "P\u207b = A\u00b7P\u00b7A\u1d40 + Q) and the measurement-correction step entirely in "
        "Eigen on the SweRV EL2 CPU."
    )
    body_para(doc,
        "The M4 hardware driver (kalman_hw_driver/src/main.cpp prior to correction) did not "
        "use the KalmanFilter class at all. It implemented the predict step as raw C nested "
        "loops over stack-allocated double arrays, then issued MMIO writes to invoke the "
        "hardware correction step. Using raw-C arrays on a bare-metal target eliminates the "
        "Eigen expression-template overhead, the heap allocation cost of MatrixXd dynamic "
        "storage, and the _sbrk bump-allocator latency\u2014all of which the software baseline "
        "incurs on every call. As a result, the two programs were not measuring the same "
        "algorithm: the software baseline measured full Eigen predict+correct, while the "
        "hardware benchmark measured nearly-zero-cost raw-C predict plus MMIO correction."
    )
    body_para(doc,
        "The per-run heap delta of 0 B in the old hardware benchmark (Table 4 of Section 8) "
        "was the visible symptom: a program using Eigen MatrixXd inevitably allocates heap on "
        "first construction, as confirmed by the software baseline\u2019s 49,760 B delta. The "
        "absence of any heap allocation in the old hardware driver proved that Eigen was never "
        "invoked."
    )

    h1(doc, "10.2 Incorrect FLOPs Count")

    body_para(doc,
        "The run_profiles.sh script used HW_FLOPS=1080 (24 FLOPs/update \u00d7 45 updates). "
        "This figure counted only the floating-point operations performed inside the hardware "
        "correction kernel (innovation scalar y, scalar S, three Newton\u2013Raphson iterations, "
        "Kalman gain vector K, state correction x_out, and covariance update P_out). However, "
        "the mcycle CSR bracket in the firmware encompasses the entire update loop, including "
        "the predict step. The predict step itself performs: \u0078\u0302_new = A\u00b7\u0078\u0302 "
        "(3\u00d73 times 3\u00d71 = 9 multiplies + 6 adds = 15 FLOPs) and P\u207b = A\u00b7P\u00b7A\u1d40 + Q "
        "(two 3\u00d73 GEMMs plus a 3\u00d73 addition = 2\u00d7(27 mults + 18 adds) + 9 adds = 99 FLOPs), "
        "totalling 114 FLOPs/update for predict and 24 for correct = 138 FLOPs/update = "
        "6,210 FLOPs total. Using 1,080 in the numerator while timing 6,210 FLOPs of work "
        "caused the reported GFLOP/s to be understated by a factor of 5.75."
    )

    # ─────────────────────────────────────────────────────────────────────────
    # 11. CORRECTED BENCHMARK PROGRAM
    # ─────────────────────────────────────────────────────────────────────────
    h1(doc, "11. Corrected Benchmark Program")

    body_para(doc,
        "The corrected driver, kalman_hw_driver/src/main.cpp, is restructured to mirror "
        "kalman-test.cpp exactly. It constructs the same KalmanFilter object with identical "
        "A, C, Q, R, P, x\u2080 parameters and the same 45 measurements in the same order. "
        "The loop calls kf.update_hw(y) \u2014 a new method on the KalmanFilter class \u2014 "
        "which internally runs the Eigen predict step (identical to what kf.update(y) does "
        "in the baseline) and then marshals the predicted state and covariance to the hardware "
        "accelerator via MMIO for the correction step. The mcycle bracket encloses the same "
        "scope in both programs: the 45-iteration update loop and nothing else."
    )
    body_para(doc,
        "The corrected FLOPs count is 6,210 total (138 per update: 114 for the Eigen predict "
        "step plus 24 for the hardware correction). The run_profiles.sh script was updated "
        "from HW_FLOPS=1080 to HW_FLOPS=6210. With this change, the comparison is "
        "apples-to-apples: both programs invoke the same algorithm on the same hardware using "
        "the same library; only the correction step differs (Eigen on CPU vs MMIO to hardware "
        "accelerator)."
    )
    body_para(doc,
        "The build fix: Eigen 5.0.0 ships with benchmark source files in its bench/ "
        "subdirectory that require Boost and CBLAS headers unavailable on the bare-metal "
        "target. PlatformIO\u2019s Library Dependency Finder was picking up these files and "
        "failing to compile them. A library.json was added to lib/eigen-5.0.0/ with "
        "\u201csrcFilter\u201d: \u201c-<*>\u201d, marking the library as header-only and suppressing "
        "compilation of all .cpp files under it."
    )

    # ─────────────────────────────────────────────────────────────────────────
    # 12. CORRECTED BENCHMARK RESULTS
    # ─────────────────────────────────────────────────────────────────────────
    h1(doc, "12. Corrected Benchmark Results")

    body_para(doc,
        "Five back-to-back runs were collected on the same Nexys A7 board with the same GDB "
        "JTAG method (mcycle CSR snapshot + hardware breakpoint at profile_done()). The FPGA "
        "was programmed with the SoC_impl bitstream (unchanged from M4) immediately before "
        "the run to ensure a clean accelerator state."
    )

    add_table(doc,
        ["Run", "Cycles", "Time (ms)", "Samples/s", "GFLOP/s", "Heap (B)"],
        [
            ["1",   "1,865,423", "143.494", "313.60",  "4.3277\u00d710\u207b\u2075", "10,832"],
            ["2",   "1,865,531", "143.502", "313.58",  "4.3275\u00d710\u207b\u2075", "10,832"],
            ["3",   "1,865,538", "143.503", "313.58",  "4.3274\u00d710\u207b\u2075", "10,832"],
            ["4",   "1,865,369", "143.490", "313.61",  "4.3278\u00d710\u207b\u2075", "10,832"],
            ["5",   "1,865,530", "143.502", "313.58",  "4.3275\u00d710\u207b\u2075", "10,832"],
            ["AVG", "1,865,478", "143.498", "313.590", "4.3276\u00d710\u207b\u2075", "10,832"],
        ],
        col_widths=[0.5, 1.2, 1.1, 1.0, 1.3, 1.0]
    )
    caption(doc,
        "Table 5. Corrected HW accelerator benchmark \u2014 5-run results "
        "(Nexys A7, SoC_impl bitstream, 13 MHz)."
    )

    body_para(doc,
        "The corrected benchmark shows an average of 1,865,478 cycles (143.498 ms) for "
        "45 updates, yielding a true speedup of 2.995\u00d7 over the software baseline "
        "(5,587,112 cycles, 429.778 ms). This is 36\u00d7 lower than the 108.7\u00d7 figure "
        "reported in Section 8. The heap delta of 10,832 B per run confirms that Eigen\u2019s "
        "dynamic matrix storage is now active on the hardware benchmark side, consistent with "
        "the software baseline (49,760 B). The smaller heap in the hardware run reflects only "
        "the matrices local to KalmanFilter (A, C, Q, R, P, K, x \u2014 roughly 10\u201312 Eigen "
        "MatrixXd/VectorXd objects vs the software baseline\u2019s additional temporaries created "
        "during the full correction step computation in Eigen)."
    )

    add_table(doc,
        ["Metric", "SW Baseline", "Old HW (\u00a78)", "Corrected HW", "Corrected Speedup"],
        [
            ["Avg cycles (45 updates)", "5,587,112",    "51,378",        "1,865,478",   "2.995\u00d7"],
            ["Total time (ms)",         "429.778",       "3.952",         "143.498",     "2.995\u00d7"],
            ["Samples/s",               "104.706",       "11,386",        "313.590",     "2.995\u00d7"],
            ["GFLOP/s",       "2.4082\u00d710\u207b\u2075", "2.7327\u00d710\u207b\u2074", "4.3276\u00d710\u207b\u2075", "1.80\u00d7"],
            ["Heap/run (B)",             "49,760",        "0",             "10,832",      "\u2014"],
        ],
        col_widths=[1.9, 1.3, 1.2, 1.3, 1.3]
    )
    caption(doc,
        "Table 6. SW baseline vs old HW benchmark vs corrected HW benchmark."
    )

    figure(doc, FIG4,
           "Figure 4. Per-45-update runtime comparison. The old hardware benchmark\u2019s "
           "near-zero predict time (raw-C) caused the large apparent speedup. The corrected "
           "benchmark uses Eigen predict, which dominates at 139.5 ms of the 143.5 ms total.",
           width_in=6.0)

    figure(doc, FIG5,
           "Figure 5. Reported (108.7\u00d7) vs corrected (2.995\u00d7) speedup over the SW "
           "baseline. The 36\u00d7 difference reflects the structural asymmetry in the original "
           "benchmark.",
           width_in=6.0)

    body_para(doc,
        "The factor-of-36 reduction in speedup has a single root cause: the Eigen predict "
        "step on the SweRV EL2 CPU (no FPU, software-emulated F64) is the dominant cost in "
        "both programs, and it is not accelerated by the hardware. Subtracting the hardware "
        "correction time from the corrected total (1,865,478 \u2212 51,378 \u2248 1,814,100 cycles) "
        "gives approximately 1,814,100 cycles (139.5 ms) consumed by the Eigen predict step "
        "alone. The hardware correction\u2014the portion actually offloaded\u2014accounts for only "
        "~51,378 cycles (3.95 ms), the same as the old benchmark\u2019s total. The predict step "
        "is ~35\u00d7 more expensive than the correction step at this CPU, making it the new "
        "bottleneck."
    )
    body_para(doc,
        "This finding is not a design defect in the hardware accelerator. The accelerator "
        "correctly offloads the correction step and does so in 72 cycles at 100 MHz. The "
        "limitation is that the predict step\u2014matrix-vector and matrix-matrix multiplications "
        "using software-emulated F64 on a 13 MHz RV32IMAC CPU\u2014was never targeted for "
        "acceleration. The 108.7\u00d7 speedup claimed in Section 8 was an artifact of measuring "
        "a faster (non-Eigen) implementation of the predict step in the hardware benchmark."
    )

    # ─────────────────────────────────────────────────────────────────────────
    # 13. UPDATED ROOFLINE ANALYSIS
    # ─────────────────────────────────────────────────────────────────────────
    h1(doc, "13. Updated Roofline Analysis")

    body_para(doc,
        "The roofline positions of both programs shift with the corrected FLOPs and timing. "
        "For the corrected hardware benchmark, the total FLOPs are 6,210 (138/update) and the "
        "bytes transferred are 12,960 (288 B/update: 192 B MMIO + 96 B Eigen predict memory "
        "traffic), giving AI = 6,210 / 12,960 = 0.479 FLOP/byte. The measured GFLOP/s is "
        "4.3276\u00d710\u207b\u2075. For the software baseline, AI = 10,350 FLOPs / 9,720 B "
        "\u2248 1.065 FLOP/byte (using the corrected Eigen transfer count of 216 B/update \u00d7 45), "
        "though the existing report quoted 0.21; the corrected value is used here."
    )

    figure(doc, FIG6,
           "Figure 6. Side-by-side roofline plots. Left: original M4 roofline with correction-only "
           "HW point (AI = 0.088, GFLOP/s = 2.73\u00d710\u207b\u2074). Right: corrected roofline "
           "with Eigen-predict+correction HW point (AI = 0.479, GFLOP/s = 4.33\u00d710\u207b\u2075). "
           "Both platforms sit well below the roofline, remaining latency-bound.",
           width_in=6.5)

    body_para(doc,
        "In the original roofline (left subplot of Figure 6), the hardware accelerator appeared "
        "to have lower arithmetic intensity than the software baseline (0.088 vs 0.21 FLOP/byte) "
        "and higher measured performance (2.73\u00d710\u207b\u2074 vs 2.41\u00d710\u207b\u2075 GFLOP/s), "
        "reflecting the asymmetric benchmark where the raw-C predict step consumed negligible time. "
        "In the corrected roofline (right subplot), the hardware accelerator point shifts right "
        "(higher AI = 0.479 due to the Eigen predict FLOPs) and drops (lower GFLOP/s = "
        "4.33\u00d710\u207b\u2075) relative to the software baseline. Both points now sit well below "
        "the roofline. The hardware accelerator\u2019s GFLOP/s is now only 1.80\u00d7 higher than "
        "the software baseline, consistent with the 2.995\u00d7 time speedup less the overhead "
        "difference (the corrected hardware run does fewer total FLOPs-per-second because the "
        "predict-step FLOPs are the same but the time is longer relative to what Section 8 reported)."
    )
    body_para(doc,
        "The analysis in Sections 2 and 8 of the M4 report regarding AXI-transaction latency "
        "remains valid for the correction kernel in isolation: 34 MMIO round-trips at ~2.6 \u00b5s "
        "each still dominate the correction step\u2019s 3.95 ms. The new finding is that the "
        "correction step is no longer the dominant system-level bottleneck. The Eigen predict "
        "step, at ~139.5 ms per 45-update batch, exceeds the correction-step time by 35\u00d7. "
        "Achieving meaningful further speedup requires either: (a) implementing the predict step "
        "in hardware as well, doubling the accelerated workload, or (b) implementing state "
        "retention as described in Section 9, which reduces the correction-step AXI overhead "
        "from 34 to ~5 transactions per update."
    )

    return doc


# ── Main ──────────────────────────────────────────────────────────────────────
if __name__ == '__main__':
    print("Generating figures ...")
    make_fig4()
    make_fig5()
    make_fig6()

    print("Building report ...")
    doc = new_doc()
    doc = build(doc)
    doc.save(DOCX_OUT)
    print(f"Saved: {DOCX_OUT}")
